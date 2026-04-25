#data_loader.py
import os
import io
import re
import asyncio
from config import settings
import nest_asyncio
import cv2
import pytesseract
import numpy as np
import pandas as pd
from langchain_core.documents import Document
from pathlib import Path
from PIL import Image
from pdf2image import convert_from_path
from docx import Document as DocxDocument
from docx.oxml.ns import qn
import hashlib
import uuid
import fitz

from llm_clients import query_ollama
import tiktoken

ENCODER = tiktoken.get_encoding("cl100k_base")
EMBED_MAX_TOKENS = 400 

def count_tokens(text: str) -> int:
    return len(ENCODER.encode(text))

#---------------------------------------------------------------------------

#  NEW: Context-Prefixed Markdown (The Token-Optimized Table Solution)
def table_to_context_markdown(
    df: pd.DataFrame,
    source: str,
    section: str,
    max_tokens: int = EMBED_MAX_TOKENS,   # ADDED: token budget parameter
    ) -> str:
    """
    Converts a dataframe chunk to a Markdown table and prepends a semantic context header.
    This solves the 'RAG Table Dilemma' by combining token-efficiency (Markdown)
    with high retrieval accuracy (Semantic Prefix mapping the columns).
    """
    valid_columns = [str(c).strip() for c in df.columns if str(c).strip() and str(c).lower() != "nan"]
    columns_text = ", ".join(valid_columns) if valid_columns else "unnamed columns"

    # SHORTENED prefix: same semantic meaning, fewer tokens wasted on filler words
    prefix = (
        f"Structured table from '{source}', section '{section}'. "
        f"Columns: {columns_text}.\n\n"
    )

    prefix_tokens = len(ENCODER.encode(prefix))   # ADDED: measure how many tokens the prefix itself costs
    budget = max_tokens - prefix_tokens            # ADDED: remaining tokens available for the actual table rows

    # ADDED: try progressively fewer rows until the markdown fits within the token budget
    sample_md = df.head(min(5, len(df))).to_markdown(index=False) or ""
    tokens_per_row = max(1, len(ENCODER.encode(sample_md)) / min(5, len(df)))
    safe_rows = max(1, int(budget / tokens_per_row))
    md = df.head(safe_rows).to_markdown(index=False)
    return prefix + md
# ---------------------------------------------------------------------------
class SemanticAwareTextSplitter:
    """
    Splits text into chunks that respect **semantic** (topic-shift) boundaries
    rather than fixed character counts.
    """
    # MIN_CHUNK_CHARS:    int   = 400
    # BASE_CHUNK_CHARS:   int   = 1_200
    # MAX_CHUNK_CHARS:    int   = 2_000
    BOUNDARY_THRESHOLD: float = 0.35   # Minimum cosine distance to be a boundary
    BRIDGE_SENTENCES:   int   = 1      # Sentences carried forward for context
    MIN_TOKENS = 200
    MAX_TOKENS = 500
    OVERLAP_TOKENS = 50
    _ABBREV_RE = re.compile(
        r'\b(Mr|Mrs|Ms|Dr|Prof|Sr|Jr|vs|etc|e\.g|i\.e|Fig|No|Vol|pp)\.\s',
        re.IGNORECASE,
    )
    _DECIMAL_RE = re.compile(r'(\d+)\.(\d+)')
    _SENT_SPLIT_RE = re.compile(r'(?<=[.!?…])\s+(?=[A-Z"\'\(])')

    def __init__(self) -> None:
        self._has_sklearn = self._try_import_sklearn()

    def _try_import_sklearn(self) -> bool:
        try:
            from sklearn.feature_extraction.text import TfidfVectorizer
            from sklearn.metrics.pairwise import cosine_similarity
            self._TfidfVectorizer  = TfidfVectorizer
            self._cosine_similarity = cosine_similarity
            return True
        except ImportError:
            return False

    # def compute_target_size(self, text: str) -> int:
    #     words = text.split()
    #     if not words:
    #         return self.BASE_CHUNK_CHARS

    #     unique_ratio = len({w.lower() for w in words}) / len(words)
    #     avg_word_len = sum(len(w) for w in words) / len(words)

    #     density = min(1.0,
    #                   unique_ratio * 0.6
    #                   + min(avg_word_len, 8) / 8 * 0.4)

    #     target = int(
    #         self.MAX_CHUNK_CHARS
    #         - density * (self.MAX_CHUNK_CHARS - self.MIN_CHUNK_CHARS)
    #     )
    #     return max(self.MIN_CHUNK_CHARS, min(self.MAX_CHUNK_CHARS, target))

    def tokenize_sentences(self, text: str) -> list[str]:
        protected = self._ABBREV_RE.sub(
            lambda m: m.group().replace(". ", ".\x00"), text
        )
        protected = self._DECIMAL_RE.sub(
            lambda m: m.group().replace(".", "\x01"), protected
        )

        raw_parts = self._SENT_SPLIT_RE.split(protected)

        sentences: list[str] = []
        for part in raw_parts:
            restored = part.replace("\x00", ". ").replace("\x01", ".").strip()
            if len(restored) >= 10:
                sentences.append(restored)

        return sentences or [text.strip()]

    # def _pairwise_distances(self, sentences: list[str]) -> list[float]:
    #     if len(sentences) < 2:
    #         return []

    #     if self._has_sklearn:
    #         try:
    #             vec = self._TfidfVectorizer(
    #                 analyzer="char_wb", ngram_range=(3, 5), min_df=1
    #             )
    #             matrix = vec.fit_transform(sentences)
    #             return [
    #                 1.0 - float(
    #                     self._cosine_similarity(matrix[i], matrix[i + 1])[0][0]
    #                 )
    #                 for i in range(len(sentences) - 1)
    #             ]
    #         except Exception:
    #             pass  # Fall through to Jaccard

    #     def _jaccard(a: str, b: str) -> float:
    #         sa, sb = set(a.lower().split()), set(b.lower().split())
    #         if not sa and not sb:
    #             return 0.0
    #         return 1.0 - len(sa & sb) / len(sa | sb)

    #     return [_jaccard(sentences[i], sentences[i + 1])
    #             for i in range(len(sentences) - 1)]

    # def find_semantic_boundaries(self, sentences: list[str]) -> set[int]:
    #     distances = self._pairwise_distances(sentences)
    #     if not distances:
    #         return set()

    #     arr = np.array(distances)
    #     adaptive_threshold = max(
    #         self.BOUNDARY_THRESHOLD,
    #         float(arr.mean() + 0.5 * arr.std()),
    #     )

    #     boundaries: set[int] = set()
    #     for i, d in enumerate(distances):
    #         if d < adaptive_threshold:
    #             continue
    #         left  = distances[i - 1] if i > 0                   else -1.0
    #         right = distances[i + 1] if i + 1 < len(distances)  else -1.0
    #         if d >= left and d >= right:
    #             boundaries.add(i + 1)

    #     return boundaries

    def split(self, text: str) -> list[str]:
        sentences = self.tokenize_sentences(text)
        if not sentences:
            return []

        chunks = []
        current = []
        current_tokens = 0

        for sent in sentences:
            sent = sent.strip()
            if not sent:
                continue

            sent_tokens = count_tokens(sent)

            # If exceeds → flush
            if current_tokens + sent_tokens > self.MAX_TOKENS:
                if current:
                    chunks.append(" ".join(current))

                    # 🔁 Token overlap
                    overlap = []
                    overlap_tokens = 0

                    for s in reversed(current):
                        t = count_tokens(s)
                        overlap.insert(0, s)
                        overlap_tokens += t
                        if overlap_tokens >= self.OVERLAP_TOKENS:
                            break

                    current = overlap
                    current_tokens = overlap_tokens
                else:
                    # single long sentence
                    chunks.append(sent)
                    continue

            current.append(sent)
            current_tokens += sent_tokens

        if current:
            chunks.append(" ".join(current))

        return chunks

    @staticmethod
    def dynamic_table_rows(df: "pd.DataFrame", target_tokens: int = 350) -> int:
        if df.empty:
            return 20  # CHANGED: 30 → 20, safer default since we're now token-aware
        sample = min(5, len(df))  # CHANGED: 10 → 5, 5 rows is enough to estimate tokens per row
        try:
            sample_md = df.head(sample).to_markdown(index=False) or ""
            tokens_per_row = max(1, len(ENCODER.encode(sample_md)) / sample)  # CHANGED: len(sample_md) → len(ENCODER.encode(sample_md)), measures actual tokens not characters
            rows = max(5, min(60, int(target_tokens / tokens_per_row)))  # CHANGED: target_chars → target_tokens, now dividing tokens by tokens per row which is consistent
        except Exception:
            rows = 20  # CHANGED: 30 → 20, matches the safer default above
        return rows


# ---------------------------------------------------------------------------
# Heading / topic-aware splitter
# ---------------------------------------------------------------------------

class HeaderAwareTextSplitter:
    CONTEXT_MODEL = settings.CONTEXT_MODEL

    def __init__(self):
        self._semantic_splitter = SemanticAwareTextSplitter()

    def _generate_chunk_id(self, text: str, meta: dict) -> str:
        source = meta.get("source", "unknown")
        unique_string = f"{source}::{text[:100]}"
        return hashlib.sha256(unique_string.encode('utf-8')).hexdigest()[:12]
    
    async def _generate_chunk_context(self, doc_title: str, section: str, chunk_text: str) -> str:
        prompt = (
            f"Document: {doc_title}\n"
            f"Section: {section or 'main body'}\n\n"
            f"Chunk:\n{chunk_text[:1000]}\n\n"
            "Write 1-2 sentences explaining what this chunk is about within the document. "
            "Be specific. Return only those sentences. Do not use conversational filler."
        )
        try:
            response_gen = query_ollama(prompt=prompt, model=self.CONTEXT_MODEL, keep_alive=0,num_ctx_tokens=10_000, stream=False)
            parts = []
            async for chunk in response_gen:
                parts.append(chunk)
            return "".join(parts).strip()
        except Exception as e:
            print(f"⚠️ Context generation failed: {e}")
            return ""

    def split_sections(
        self,
        sections: list[tuple[str, str]],
        base_metadata: dict,
        is_table: bool = False,
        ) -> list[Document]:
        docs = []

        for header, body in sections:
            body = body.strip()
            if not body:
                continue

            # ✅ metadata only (no content pollution)
            meta = base_metadata.copy()
            if header:
                meta["section"] = header
            meta["is_table"] = is_table

            # -------------------------
            # ✅ TABLES (no prefix noise)
            # -------------------------
            if is_table:
                meta["chunk_id"] = self._generate_chunk_id(body, meta)

                docs.append(
                    Document(
                        page_content=body,   # ✅ clean table content only
                        metadata=meta
                    )
                )
                continue

            # -------------------------
            # ✅ TEXT SPLITTING (token-based only)
            # -------------------------
            sub_texts = self._semantic_splitter.split(body)

            for idx, sub_text in enumerate(sub_texts):
                sub_text = sub_text.strip()
                if not sub_text:
                    continue

                sub_meta = meta.copy()
                sub_meta["chunk_id"] = f"{self._generate_chunk_id(sub_text, meta)}_{idx}"
                sub_meta["part"] = f"{idx + 1}/{len(sub_texts)}"

                docs.append(
                    Document(
                        page_content=sub_text,   # ✅ NO HEADER, NO LABEL
                        metadata=sub_meta
                    )
                )

        return docs
    
    def split_sections_with_context(
        self,
        sections: list[tuple[str, str]],
        base_metadata: dict,
        is_table: bool = False,
    ) -> list[Document]:
        """
        Splits sections and adds context. Tables bypass the LLM generation 
        because their semantic prefix is already mapped during markdown conversion.
        """
        docs = self.split_sections(sections, base_metadata, is_table=is_table)
        doc_title = base_metadata.get("source", "unknown_document")

        if is_table:
            # 🔥 tables already have their explicit context mapped by table_to_context_markdown
            for doc in docs:
                doc.metadata["embed_content"] = doc.page_content
            return docs

        sem = asyncio.Semaphore(2)

        async def bounded_generate(d: Document) -> str:
            async with sem:
                return await self._generate_chunk_context(
                    doc_title,
                    d.metadata.get("section", ""),
                    d.page_content
                )

        async def enrich_all():
            tasks = [bounded_generate(d) for d in docs]
            return await asyncio.gather(*tasks)

        nest_asyncio.apply()          # patch BEFORE any event loop interaction
        contexts = asyncio.run(enrich_all())

        for doc, ctx in zip(docs, contexts):
            doc.metadata["context_prefix"] = ctx
            doc.metadata["embed_content"] = doc.page_content

        return docs

        
    
    def split_plain_text(self, text: str, base_metadata: dict) -> list[Document]:
        sections = _detect_text_headers(text)
        if sections:
            return self.split_sections(sections, base_metadata)
        sections = _paragraph_sections(text=text, splitter=self._semantic_splitter)
        return self.split_sections(sections, base_metadata)


# ---------------------------------------------------------------------------
# Header detection helpers
# ---------------------------------------------------------------------------

_MARKDOWN_HEADER_RE = re.compile(r"^(#{1,4})\s+(.+)$", re.MULTILINE)
_GENERIC_HEADER_RE = re.compile(
    r"^(?:"
    r"(?:\d+[.\d]+\s+[A-Z].{3,60})" 
    r"|(?:[A-Z0-9][A-Z0-9\s\-&]{4,60}[A-Z0-9])" 
    r"|(?:(?:Chapter|Section|Part|Appendix)\s*[\dIVXivx]+.?)" 
    r")$",
    re.MULTILINE | re.IGNORECASE,
)

def _detect_text_headers(text: str) -> list[tuple[str, str]]:
    matches = list(_MARKDOWN_HEADER_RE.finditer(text))
    if len(matches) >= 2:
        return _build_sections_from_matches(text, matches, label_group=2)

    matches = list(_GENERIC_HEADER_RE.finditer(text))
    if len(matches) >= 2:
        return _build_sections_from_matches(text, matches, label_group=0)

    return []

def _build_sections_from_matches(text: str, matches: list, label_group: int) -> list[tuple[str, str]]:
    sections = []
    for idx, match in enumerate(matches):
        header = match.group(label_group).strip()
        body_start = match.end()
        body_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        body = text[body_start:body_end].strip()
        sections.append((header, body))
    return sections

def _paragraph_sections(text: str, splitter: "SemanticAwareTextSplitter | None" =None, min_chars: int = 100) -> list[tuple[str, str]]:
    raw_paras = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    merged: list[str] = []
    buffer = ""
    for para in raw_paras:
        buffer = (buffer + "\n\n" + para).strip() if buffer else para
        if len(buffer) >= min_chars:
            merged.append(buffer)
            buffer = ""
    if buffer:
        if merged:
            merged[-1] += "\n\n" + buffer
        else:
            merged.append(buffer)

    if not splitter:
        return [("", block) for block in merged]

    sections: list[tuple[str, str]] = []
    max_size = splitter.MAX_TOKENS
    for block in merged:
        if count_tokens(block) > max_size:
            for sub in splitter.split(block):
                if sub.strip():
                    sections.append(("", sub.strip()))
        else:
            sections.append(("", block))
    return sections


# ---------------------------------------------------------------------------
# DOCX-specific section extractor
# ---------------------------------------------------------------------------
def _is_toc_page(text: str) -> bool:
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if not lines:
        return False
    toc_line_pattern = re.compile(r'\.{4,}\s*\d+$')
    toc_count = sum(1 for line in lines if toc_line_pattern.search(line))
    ratio = toc_count / len(lines)
    return ratio > 0.25

def _extract_docx_sections(docx: DocxDocument, source_name: str) -> list[tuple[str, str]]:
    text_sections: list[tuple[str, str]] = []
    table_sections: list[tuple[str, str]] = []

    current_heading = ""
    current_body_lines: list[str] = []

    for para in docx.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            continue

        if style_name.startswith("Heading"):
            if current_body_lines:
                text_sections.append((current_heading, "\n".join(current_body_lines)))
            current_heading = text
            current_body_lines = []
        else:
            current_body_lines.append(text)

    if current_body_lines:
        text_sections.append((current_heading, "\n".join(current_body_lines)))

    _semantic_splitter = SemanticAwareTextSplitter()

    for i, table in enumerate(docx.tables):
        data = []
        for row in table.rows:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            data.append(row_data)

        if len(data) > 1:
            df = pd.DataFrame(data[1:], columns=data[0])

            try:
                rows_per_chunk = max(20, _semantic_splitter.dynamic_table_rows(df))
                for start_row in range(0, len(df), rows_per_chunk):
                    df_chunk = df.iloc[start_row:start_row + rows_per_chunk]

                    #  NEW: Use Context-Prefixed Markdown
                    content = table_to_context_markdown(
                        df_chunk, 
                        source_name, 
                        current_heading or f"Table {i + 1}"
                    )

                    if content:
                        end_row = start_row + len(df_chunk)
                        header = f"{current_heading or 'DOCX Section'} — Table {i + 1} (rows {start_row + 1}–{end_row})"
                        table_sections.append((header, content))

            except Exception as e:
                print(f"Skipping malformed DOCX table: {e}")

    return text_sections, table_sections


# ---------------------------------------------------------------------------
# PDF-specific header heuristic
# ---------------------------------------------------------------------------

def _extract_pdf_sections(text: str) -> list[tuple[str, str]]:
    return _detect_text_headers(text) or _paragraph_sections(text)


# ---------------------------------------------------------------------------
# Main loader class
# ---------------------------------------------------------------------------
class MultiFormatDocumentLoader:
    def __init__(self):
        self.supported_extensions = {
            ".pdf", ".docx", ".csv", ".xlsx", ".xls", ".txt",
            ".png", ".jpg", ".jpeg", ".bmp", ".tiff"
        }
        self.splitter = HeaderAwareTextSplitter()
        self._semantic_splitter = SemanticAwareTextSplitter()

    def _detect_table_grid_cv2(self, image_path: str) -> bool:
        try:
            img = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
            if img is None:
                return False
            _, thresh = cv2.threshold(img, 128, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)
            line_min_length = np.array(img).shape[1] // 10 
            kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (line_min_length, 1))
            kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, line_min_length))
            horizontal = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel_h)
            vertical = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel_v)
            intersections = cv2.bitwise_and(horizontal, vertical)
            contours, _ = cv2.findContours(intersections, cv2.RETR_CCOMP, cv2.CHAIN_APPROX_SIMPLE)
            return len(contours) > 6
        except Exception as e:
            print(f"⚠️ Error during OpenCV table detection: {e}")
            return False

    def _looks_like_borderless_table(self, text: str) -> bool:
        if not text:
            return False
        lines = text.split("\n")
        tabular_line_count = sum(
            1 for line in lines if len(re.split(r"\s{3,}", line.strip())) >= 3
        )
        return tabular_line_count >= 3

    async def _get_ocr_text_async(self, file_path: str, origin_info: str = "Image") -> str:
        ocr_text = ""
        confidence = 0
        has_grid = await asyncio.to_thread(self._detect_table_grid_cv2, file_path)
        if not has_grid:
            try:
                def run_tesseract():
                    img = cv2.imread(file_path)
                    if img is not None:
                        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        ocr_data = pytesseract.image_to_data(thresh, output_type=pytesseract.Output.DICT)
                        ocr_text = pytesseract.image_to_string(thresh).strip()
                        confidences = [int(c) for c in ocr_data["conf"] if int(c) != -1]
                        confidence = sum(confidences) / len(confidences) if confidences else 0
                        return ocr_text, confidence
                ocr_text, confidence = await asyncio.to_thread(run_tesseract)
            except Exception as e:
                print(f"⚠️  Tesseract failed for {origin_info}: {e}")

        if has_grid or self._should_use_vlm_fallback(ocr_text, confidence):
            reason = "Grid detected" if has_grid else "Low quality/Borderless table detected"
            print(f"🔄 {reason} for {origin_info}. Falling back to VLM...")
            try:
                vlm_prompt = (
                    "Extract all readable text from this image. "
                    "CRITICAL INSTRUCTION: If the image contains a table, grid, "
                    "or structured data, you MUST format it as a valid Markdown "
                    "table using '|' to separate columns and '-' to separate the "
                    "header row. Preserve the exact row and column structure. "
                    "Return ONLY the markdown content, no conversational filler."
                )
                parts = []
                async for chunk in query_ollama(
                    prompt=vlm_prompt + origin_info,
                    model="codez-ocr:latest",
                    image_path=file_path,
                    keep_alive=0,
                    stream=False,
                ):
                    parts.append(chunk)
                ocr_text = "".join(parts).strip()
            except Exception as e:
                print(f"❌ VLM fallback failed: {e}")

        return ocr_text

    def _should_use_vlm_fallback(self, text: str, confidence: float) -> bool:
        if not text or len(text) < 10 or confidence < 60:
            return True
        if self._looks_like_borderless_table(text):
            return True
        words = text.split()
        if not words:
            return True
        return (sum(len(w) for w in words) / len(words)) < 2.5

    def _run_async_ocr(self, file_path: str, origin_info: str) -> str:
        nest_asyncio.apply()  # only applied when this specific function runs
        return asyncio.run(self._get_ocr_text_async(file_path, origin_info))

    def load_image(self, file_path: str) -> list[Document]:
        print(f"⏳ Processing image: {os.path.basename(file_path)}")
        text = self._run_async_ocr(file_path, os.path.basename(file_path))
        if not text:
            return []
        metadata = {"source": os.path.basename(file_path), "file_type": "image"}
        return self.splitter.split_plain_text(text, metadata)

    def load_pdf(self, file_path: str) -> list[Document]:
        if os.path.getsize(file_path) == 0:
            print(f"Skipping empty file: {file_path}")
            return []

        table_docs: list[Document] = []
        page_texts: list[tuple[int, str]] = []

        base_metadata = {
            "source": os.path.basename(file_path),
            "file_type": "pdf",
        }

        try:
            with fitz.open(file_path) as f:
                for page_num, page in enumerate(f):
                    page_text = page.get_text().strip() or ""

                    if _is_toc_page(page_text):
                        print(f"⏩ Skipping TOC on page {page_num + 1}")
                        continue
                    print(f"📄 Processing page {page_num + 1}")

                    try:
                        tables = page.find_tables()
                    except Exception as e:
                        print(f"⚠️ Table extraction failed on page {page_num + 1}: {e}")
                        tables = []

                    if tables:
                        page_text += "\n\n### Extracted Tabular Data\n\n"

                        for idx, tab in enumerate(tables):
                            try:
                                df = tab.to_pandas()
                                df.dropna(how="all", inplace=True)
                                
                                #  SAFETY: Skip huge tables
                                if len(df) > 2000:
                                    print(f"⚠️ Skipping huge table on page {page_num + 1}")
                                    continue

                                if not df.empty:
                                    rows_per_chunk = max(
                                        20,
                                        self._semantic_splitter.dynamic_table_rows(df)
                                    )

                                    for start_row in range(0, len(df), rows_per_chunk):
                                        df_chunk = df.iloc[start_row:start_row + rows_per_chunk]

                                        # 🔥 NEW: Context-Prefixed Markdown
                                        content = table_to_context_markdown(
                                            df_chunk,
                                            base_metadata["source"],
                                            f"Table {idx + 1} (page {page_num + 1})"
                                        )

                                        if content:
                                            end_row = start_row + len(df_chunk)
                                            table_meta = {
                                                **base_metadata,
                                                "page": page_num + 1,
                                                "data_type": "table",
                                            }
                                            header = f"Table {idx + 1} (page {page_num + 1}, rows {start_row + 1}–{end_row})"

                                            table_docs.extend(
                                                self.splitter.split_sections_with_context(
                                                    [(header, content)],
                                                    table_meta,
                                                    is_table=True,
                                                )
                                            )
                            except Exception as e:
                                print(f"Error extracting table from PDF page: {e}")

                    if page_text:
                        page_texts.append((page_num + 1, page_text))

        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
            return []

        text_docs: list[Document] = []
        total_text = "".join(t for _, t in page_texts)

        if len(total_text.strip()) > 200:
            bridge = ""
            for page_num, page_text in page_texts:
                text_with_bridge = (bridge + " " + page_text).strip() if bridge else page_text
                page_meta = {**base_metadata, "page": page_num}
                sections = _extract_pdf_sections(text_with_bridge)

                text_docs.extend(
                    self.splitter.split_sections_with_context(
                        sections,
                        page_meta,
                        is_table=False
                    )
                )

                sentences = page_text.split(". ")
                bridge = sentences[-1].strip() if sentences else ""

        else:
            print(f"⚠️  Scanned PDF detected: {os.path.basename(file_path)}")
            try:
                images = convert_from_path(file_path)
                for i, img in enumerate(images):
                    temp_path = f"temp_{os.getpid()}_{i}.jpg"
                    img = img.convert("RGB")
                    img.save(temp_path, "JPEG", quality=90)
                    ocr_text = self._run_async_ocr(temp_path, f"Page {i + 1}")
                    if os.path.exists(temp_path):
                        os.remove(temp_path)

                    if not ocr_text:
                        continue

                    page_meta = {
                        **base_metadata,
                        "file_type": "pdf_scanned",
                        "page": i + 1
                    }
                    sections = _detect_text_headers(ocr_text) or [("", ocr_text)]
                    text_docs.extend(
                        self.splitter.split_sections_with_context(sections, page_meta)
                    )

            except Exception as e:
                print(f"Error loading scanned PDF: {e}")
                return []

        final_docs = table_docs + text_docs

        if not final_docs and total_text.strip():
            return self.splitter.split_plain_text(total_text, base_metadata)

        return final_docs

    def load_docx(self, file_path: str) -> list[Document]:
        try:
            docx = DocxDocument(file_path)
        except Exception as e:
            print(f"Error loading DOCX {file_path}: {e}")
            return []

        source_name = os.path.basename(file_path)
        text_sections, table_sections = _extract_docx_sections(docx, source_name)
        metadata = {"source": source_name, "file_type": "docx"}
        
        docs = self.splitter.split_sections_with_context(text_sections, metadata, is_table=False)
        docs.extend(self.splitter.split_sections_with_context(table_sections, metadata, is_table=True))
        return docs

    def load_csv(self, file_path: str, rows_per_chunk: int = 250) -> list[Document]:
        docs = []
        source_name = os.path.basename(file_path)
        try:
            for i, df_chunk in enumerate(
                pd.read_csv(file_path, on_bad_lines="skip", chunksize=rows_per_chunk, low_memory=True)
            ):
                content = table_to_context_markdown(df_chunk, source_name, "CSV Data")
                if content:
                    start = i * rows_per_chunk + 1
                    docs.append(Document(
                        page_content=content,
                        metadata={
                            "source": source_name,
                            "rows": f"{start}-{start + len(df_chunk) - 1}",
                            "file_type": "csv_chunk",
                            "embed_content": content # Pre-mapped for tables
                        },
                    ))
        except Exception as e:
            print(f"Error loading CSV {file_path}: {e}")
        return docs

    def load_excel(self, file_path: str, rows_per_chunk: int = 250) -> list[Document]:
        docs = []
        source_name = os.path.basename(file_path)
        try:
            xls = pd.ExcelFile(file_path)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                for start_row in range(0, len(df), rows_per_chunk):
                    df_chunk = df.iloc[start_row : start_row + rows_per_chunk]
                    content = table_to_context_markdown(df_chunk, source_name, f"Sheet: {sheet_name}")
                    
                    if content:
                        start_idx = start_row + 1
                        end_idx = start_row + len(df_chunk)
                        docs.append(Document(
                            page_content=content,
                            metadata={
                                "source": source_name,
                                "sheet": sheet_name,
                                "rows": f"{start_idx}-{end_idx}",
                                "file_type": "excel_chunk",
                                "is_table": True,
                                "embed_content": content # Pre-mapped for tables
                            },
                        ))
        except Exception as e:
            print(f"Error loading Excel {file_path}: {e}")
        return docs

    def load_txt(self, file_path: str) -> list[Document]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except Exception as e:
            print(f"Error loading TXT {file_path}: {e}")
            return []

        metadata = {"source": os.path.basename(file_path), "file_type": "txt"}
        return self.splitter.split_plain_text(text, metadata)

    def load_document(self, file_path: str) -> list[Document]:
        ext = Path(file_path).suffix.lower()
        dispatch = {
            ".pdf":  self.load_pdf,
            ".docx": self.load_docx,
            ".csv":  self.load_csv,
            ".xlsx": self.load_excel,
            ".xls":  self.load_excel,
            ".txt":  self.load_txt,
            ".png":  self.load_image,
            ".jpg":  self.load_image,
            ".jpeg": self.load_image,
            ".bmp":  self.load_image,
            ".tiff": self.load_image,
        }
        loader = dispatch.get(ext)
        if loader:
            return loader(file_path)
        print(f"Unsupported file type: {ext}")
        return []

def dump_chunks_to_file(
    docs: list,
    output_path: str = "vector.txt",
    encoding: str = "utf-8",
    no_change: bool = False,
    mode: str = "w" 
) -> None:
    if no_change or not docs:
        return

    KNOWN_KEYS = ["chunk_id", "source", "file_type", "section", "page", "rows", "sheet"]
    LABELS = {
        "chunk_id":  "CHUNK ID ",
        "source":    "SOURCE   ",
        "file_type": "FILE TYPE",
        "section":   "SECTION  ",
        "page":      "PAGE     ",
        "rows":      "ROWS     ",
        "sheet":     "SHEET    ",
    }

    WIDE  = "=" * 80
    THIN  = "-" * 80

    with open(output_path, mode, encoding=encoding, errors="replace") as f:
        if mode == "w":
            f.write("=== VECTOR CHUNK DUMP ===\n\n")

        for doc in docs:
            meta = doc.metadata or {}
            content = doc.page_content or ""

            f.write(WIDE + "\n")
            f.write("CHUNK\n")
            f.write(THIN + "\n")

            for key in KNOWN_KEYS:
                if key in meta:
                    label = LABELS[key]
                    f.write(f"{label}: {meta[key]}\n")

            f.write(f"CHARS    : {len(content)}\n")

            extra = {k: v for k, v in meta.items() if k not in KNOWN_KEYS}
            if extra:
                f.write(f"OTHER    : {extra}\n")

            f.write(THIN + "\n")
            f.write(content)
            f.write("\n" + WIDE + "\n\n")