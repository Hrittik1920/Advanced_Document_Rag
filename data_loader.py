#data_loader.py
import os
import io
import re
import asyncio
from config import settings
import nest_asyncio
import cv2
import pytesseract
import numpy as np
import pandas as pd
from langchain_core.documents import Document
from pathlib import Path
from PIL import Image
from pdf2image import convert_from_path
from docx import Document as DocxDocument
from docx.oxml.ns import qn
import hashlib
import uuid
import fitz

from llm_clients import query_ollama

nest_asyncio.apply()

# ---------------------------------------------------------------------------
class SemanticAwareTextSplitter:
    """
    Splits text into chunks that respect **semantic** (topic-shift) boundaries
    rather than fixed character counts.

    Algorithm
    ---------
    1. Tokenize text into individual sentences (regex, no heavy deps).
    2. Compute a *content-density score* from type-token ratio and average
       word length → derive a target chunk size dynamically per document.
    3. Detect semantic breakpoints via sliding-window cosine distance on
       TF-IDF character n-gram vectors (falls back to Jaccard if sklearn is
       absent).  A position is a boundary only when its distance is a local
       maximum **and** exceeds an adaptive threshold (mean + 0.5 σ).
    4. Greedily accumulate sentences; flush into a chunk when a semantic
       boundary is hit (and accumulated size ≥ MIN) or the hard MAX is
       reached.
    5. Carry the final BRIDGE_SENTENCES of each chunk into the next one so
       cross-boundary context is never lost.

    Sizing contract
    ---------------
    Dense technical text  → smaller chunks  (≈ MIN_CHUNK_CHARS)
    Narrative / sparse    → larger chunks   (≈ MAX_CHUNK_CHARS)
    """

    MIN_CHUNK_CHARS:    int   = 400
    BASE_CHUNK_CHARS:   int   = 1_200
    MAX_CHUNK_CHARS:    int   = 2_000
    BOUNDARY_THRESHOLD: float = 0.35   # Minimum cosine distance to be a boundary
    BRIDGE_SENTENCES:   int   = 1      # Sentences carried forward for context

    # Regex fragments used in sentence protection
    _ABBREV_RE = re.compile(
        r'\b(Mr|Mrs|Ms|Dr|Prof|Sr|Jr|vs|etc|e\.g|i\.e|Fig|No|Vol|pp)\.\s',
        re.IGNORECASE,
    )
    _DECIMAL_RE = re.compile(r'(\d+)\.(\d+)')
    _SENT_SPLIT_RE = re.compile(r'(?<=[.!?…])\s+(?=[A-Z"\'\(])')

    def __init__(self) -> None:
        self._has_sklearn = self._try_import_sklearn()

    # ------------------------------------------------------------------
    # sklearn / fallback import
    # ------------------------------------------------------------------

    def _try_import_sklearn(self) -> bool:
        try:
            from sklearn.feature_extraction.text import TfidfVectorizer
            from sklearn.metrics.pairwise import cosine_similarity
            self._TfidfVectorizer  = TfidfVectorizer
            self._cosine_similarity = cosine_similarity
            return True
        except ImportError:
            return False

    # ------------------------------------------------------------------
    # Dynamic sizing
    # ------------------------------------------------------------------

    def compute_target_size(self, text: str) -> int:
        """
        Returns a target chunk size (chars) suited to the text's density.

        Density is a composite of:
          - Type-token ratio   (high → many unique words → dense)
          - Average word length (high → technical vocabulary)

        density ∈ [0, 1]:  0 → narrative → MAX_CHUNK_CHARS
                           1 → dense     → MIN_CHUNK_CHARS
        """
        words = text.split()
        if not words:
            return self.BASE_CHUNK_CHARS

        unique_ratio = len({w.lower() for w in words}) / len(words)
        avg_word_len = sum(len(w) for w in words) / len(words)

        # Clamp avg_word_len contribution at 8 chars (above that it's noise)
        density = min(1.0,
                      unique_ratio * 0.6
                      + min(avg_word_len, 8) / 8 * 0.4)

        target = int(
            self.MAX_CHUNK_CHARS
            - density * (self.MAX_CHUNK_CHARS - self.MIN_CHUNK_CHARS)
        )
        return max(self.MIN_CHUNK_CHARS, min(self.MAX_CHUNK_CHARS, target))

    # ------------------------------------------------------------------
    # Sentence tokenisation
    # ------------------------------------------------------------------

    def tokenize_sentences(self, text: str) -> list[str]:
        """
        Sentence-boundary detection via punctuation-aware regex.

        Protects:
          - Common abbreviations  (Dr., Fig., etc.)
          - Decimal numbers       (3.14)
          - Ellipses              (…)

        Minimum sentence length is 10 chars to suppress artefacts.
        """
        # Temporarily replace protected sequences with placeholder bytes
        protected = self._ABBREV_RE.sub(
            lambda m: m.group().replace(". ", ".\x00"), text
        )
        protected = self._DECIMAL_RE.sub(
            lambda m: m.group().replace(".", "\x01"), protected
        )

        raw_parts = self._SENT_SPLIT_RE.split(protected)

        sentences: list[str] = []
        for part in raw_parts:
            restored = part.replace("\x00", ". ").replace("\x01", ".").strip()
            if len(restored) >= 10:
                sentences.append(restored)

        return sentences or [text.strip()]

    # ------------------------------------------------------------------
    # Semantic boundary detection
    # ------------------------------------------------------------------

    def _pairwise_distances(self, sentences: list[str]) -> list[float]:
        """
        Returns cosine distance between every adjacent sentence pair.
        Falls back to Jaccard word-set distance when sklearn is absent.
        """
        if len(sentences) < 2:
            return []

        if self._has_sklearn:
            try:
                vec = self._TfidfVectorizer(
                    analyzer="char_wb", ngram_range=(3, 5), min_df=1
                )
                matrix = vec.fit_transform(sentences)
                return [
                    1.0 - float(
                        self._cosine_similarity(matrix[i], matrix[i + 1])[0][0]
                    )
                    for i in range(len(sentences) - 1)
                ]
            except Exception:
                pass  # Fall through to Jaccard

        # Jaccard fallback
        def _jaccard(a: str, b: str) -> float:
            sa, sb = set(a.lower().split()), set(b.lower().split())
            if not sa and not sb:
                return 0.0
            return 1.0 - len(sa & sb) / len(sa | sb)

        return [_jaccard(sentences[i], sentences[i + 1])
                for i in range(len(sentences) - 1)]

    def find_semantic_boundaries(self, sentences: list[str]) -> set[int]:
        """
        Returns the set of sentence indices where a new chunk should begin.

        A position qualifies when:
          - Its cosine distance is a local maximum among its neighbours, AND
          - It exceeds max(BOUNDARY_THRESHOLD, mean + 0.5 σ) of all distances.
        """
        distances = self._pairwise_distances(sentences)
        if not distances:
            return set()

        arr = np.array(distances)
        adaptive_threshold = max(
            self.BOUNDARY_THRESHOLD,
            float(arr.mean() + 0.5 * arr.std()),
        )

        boundaries: set[int] = set()
        for i, d in enumerate(distances):
            if d < adaptive_threshold:
                continue
            left  = distances[i - 1] if i > 0                   else -1.0
            right = distances[i + 1] if i + 1 < len(distances)  else -1.0
            if d >= left and d >= right:
                boundaries.add(i + 1)   # boundary starts *before* sentence[i+1]

        return boundaries

    # ------------------------------------------------------------------
    # Public split interface
    # ------------------------------------------------------------------

    def split(self, text: str) -> list[str]:
        """
        Split *text* into semantically coherent, density-sized chunks.
        Each chunk is a single string (sentences joined by spaces).
        """
        sentences = self.tokenize_sentences(text)
        if len(sentences) <= 1:
            return [text.strip()] if text.strip() else []

        target_size  = self.compute_target_size(text)
        boundaries   = self.find_semantic_boundaries(sentences)

        chunks:      list[str]  = []
        current:     list[str]  = []
        current_len: int        = 0
        bridge:      list[str]  = []

        for idx, sent in enumerate(sentences):
            sent_len   = len(sent)
            at_boundary = idx in boundaries
            over_min    = current_len >= self.MIN_CHUNK_CHARS
            would_exceed = (current_len + sent_len) > self.MAX_CHUNK_CHARS

            flush = current and (
                (at_boundary and over_min) or would_exceed
            )

            if flush:
                chunks.append(" ".join(current))
                # Bridge: carry trailing sentences for cross-boundary context
                bridge  = current[-self.BRIDGE_SENTENCES:] if self.BRIDGE_SENTENCES else []
                current = list(bridge)
                current_len = sum(len(s) for s in current)

            current.append(sent)
            current_len += sent_len

        if current:
            chunks.append(" ".join(current))

        return chunks

    # ------------------------------------------------------------------
    # Dynamic table sizing helper (used by HeaderAwareTextSplitter)
    # ------------------------------------------------------------------

    @staticmethod
    def dynamic_table_rows(df: "pd.DataFrame", target_chars: int = 1_000) -> int:
        """
        Compute how many rows to include per table chunk so that each chunk
        is approximately *target_chars* characters of Markdown.

        Bounds:  [5, 60] rows.
        """
        if df.empty:
            return 30
        sample = min(10, len(df))
        try:
            sample_md  = df.head(sample).to_markdown(index=False) or ""
            chars_per_row = max(1, len(sample_md) / sample)
            rows = max(5, min(60, int(target_chars / chars_per_row)))
        except Exception:
            rows = 30
        return rows


# ---------------------------------------------------------------------------
# Heading / topic-aware splitter  (uses SemanticAwareTextSplitter internally)
# ---------------------------------------------------------------------------

class HeaderAwareTextSplitter:
    """
    Splits text into chunks that respect section boundaries (headers/topics).

    Strategy
    --------
    1. Use format-specific section extractor to get (header, body) pairs.
    2. Bodies that fit within *target_size* chars are kept as-is.
    3. Bodies that exceed *target_size* are sub-split by SemanticAwareTextSplitter
       (sentence-level, adaptive size, semantic boundary detection).
    4. Sub-chunks carry the final sentence of the previous sub-chunk as a
       "bridge" prefix — a genuine context window, not an arbitrary overlap.
    5. With no headers, fall back to paragraph grouping then semantic splitting.
    """

    CONTEXT_MODEL = settings.CONTEXT_MODEL

    def __init__(self):
        self._semantic_splitter = SemanticAwareTextSplitter()
    def _generate_chunk_id(self, text: str, meta: dict) -> str:
        """Generates a deterministic ID based on source and content."""
        source = meta.get("source", "unknown")
        # Hash the source and the first 100 chars of content
        unique_string = f"{source}::{text[:100]}"
        return hashlib.sha256(unique_string.encode('utf-8')).hexdigest()[:12]
    
    #Adding contextual ware embed chunk generation using the ollama model
    async def _generate_chunk_context(self, doc_title: str, section: str, chunk_text: str) -> str:
        """Generate a 1-2 sentence context prefix that situates this chunk in the document."""
        prompt = (
            f"Document: {doc_title}\n"
            f"Section: {section or 'main body'}\n\n"
            f"Chunk:\n{chunk_text[:1000]}\n\n"
            "Write 1-2 sentences explaining what this chunk is about within the document. "
            "Be specific. Return only those sentences. Do not use conversational filler."
        )
        try:
            # Note: Assuming query_ollama yields text chunks based on your OCR code
            response_gen = query_ollama(prompt=prompt, model=self.CONTEXT_MODEL, keep_alive=0,num_ctx_tokens=10_000, stream=False)
            parts = []
            async for chunk in response_gen:
                parts.append(chunk)
            return "".join(parts).strip()
        except Exception as e:
            print(f"⚠️ Context generation failed: {e}")
            return ""
    
    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def split_sections(
        self,
        sections: list[tuple[str, str]],   # [(header, body), ...]
        base_metadata: dict,
        is_table: bool = False,
    ) -> list[Document]:
        """
        Convert pre-parsed (header, body) pairs into Documents.
        Applies sub-splitting only when a section body exceeds MAX_SECTION_CHARS.
        """
        docs = []
        for header, body in sections:
            body = body.strip()
            if not body:
                continue

            meta = {**base_metadata, "section": header} if header else base_metadata.copy()
            meta["is_table"] = is_table
            contextualized_body = f"[{header}]\n{body}" if header else body
            if is_table :
                meta["chunk_id"] = self._generate_chunk_id(body, meta)
                docs.append(Document(page_content=contextualized_body, metadata=meta))
                continue
            target_size = self._semantic_splitter.compute_target_size(body)

            if len(body) <= target_size:
                # Body fits in one chunk
                meta["chunk_id"] = self._generate_chunk_id(body, meta)
                docs.append(Document(page_content=contextualized_body, metadata=meta))
            else:
                # Semantically sub-split the oversized body
                sub_texts = self._semantic_splitter.split(body)

                for idx, sub_text in enumerate(sub_texts):
                    sub_text = sub_text.strip()
                    if not sub_text:
                        continue

                    is_continuation = idx > 0
                    if is_continuation and header:
                        # Prefix with bridge context already embedded by the
                        # SemanticAwareTextSplitter; add a readable label too.
                        label   = f"[{header} — part {idx + 1} of {len(sub_texts)}]"
                    elif header:
                        label = f"[{header}]"
                    else:
                        label = None

                    content = f"{label}\n{sub_text}" if label else sub_text
                    sub_meta = {**meta}
                    sub_meta["chunk_id"] = (
                        f"{self._generate_chunk_id(sub_text, meta)}_{idx}"
                    )
                    sub_meta["part"] = f"{idx + 1}/{len(sub_texts)}"
                    docs.append(
                        Document(page_content=content, metadata=sub_meta)
                    )

        return docs
    
    def split_sections_with_context(
        self,
        sections: list[tuple[str, str]],
        base_metadata: dict,
        is_table: bool = False,
    ) -> list[Document]:
        """
        Splits sections and adds an AI-generated context prefix to text chunks.
        Tables bypass context generation to save massive amounts of compute time.
        """
        # First, get the standard docs using your existing logic
        docs = self.split_sections(sections, base_metadata, is_table=is_table)

        doc_title = base_metadata.get("source", "unknown_document")

        if is_table:
            section = base_metadata.get("section", "Table")
            
            for doc in docs:
                # Manually inject a rigid context prefix for tables
                doc.metadata["embed_content"] = (f"[Source: {doc_title} | Section: {section}]\n\n{doc.page_content}")
            return docs

        # Create a Semaphore to limit Ollama to processing 4 chunks at a time
        # Change this number based on your GPU VRAM (lower if it crashes, higher if GPU usage is low)
        sem = asyncio.Semaphore(2)

        async def bounded_generate(d: Document) -> str:
            async with sem:
                return await self._generate_chunk_context(
                    doc_title,
                    d.metadata.get("section", ""),
                    d.page_content
                )

        async def enrich_all():
            tasks = [bounded_generate(d) for d in docs]
            return await asyncio.gather(*tasks)

        # Run the async generation
        try:
            loop = asyncio.get_event_loop()
            contexts = loop.run_until_complete(enrich_all())
        except RuntimeError:
            contexts = asyncio.run(enrich_all())

        for doc, ctx in zip(docs, contexts):
            doc.metadata["context_prefix"] = ctx
            doc.metadata["embed_content"] = (
                f"[{doc_title} Context: {ctx}]\n\n{doc.page_content}"
                if ctx else doc.page_content
            )

        return docs
    
    def split_plain_text(self, text: str, base_metadata: dict) -> list[Document]:
        """
        Detect headers in raw text, then call split_sections.
        Falls back to paragraph grouping if no headers are found.
        """
        sections = _detect_text_headers(text)
        if sections:
            return self.split_sections(sections, base_metadata)

        # No headers — split by paragraphs, merge small ones
        sections = _paragraph_sections(text=text, splitter=self._semantic_splitter)
        return self.split_sections(sections, base_metadata)


# ---------------------------------------------------------------------------
# Header detection helpers (module-level, no state needed)
# ---------------------------------------------------------------------------

# Matches: "# Header", "## Sub", "### Sub-sub"
_MARKDOWN_HEADER_RE = re.compile(r"^(#{1,4})\s+(.+)$", re.MULTILINE)

# Matches: "1.2.3  Some Title", "CHAPTER 3 — INTRO", all-caps short lines
_GENERIC_HEADER_RE = re.compile(
    r"^(?:"
    # Numbered: "1.2 Section Title" (Allows for extra spaces from OCR)
    r"(?:\d+[.\d]+\s+[A-Z].{3,60})" 
    # ALL CAPS line (Allows numbers and symbols like '&' often found in headers)
    r"|(?:[A-Z0-9][A-Z0-9\s\-&]{4,60}[A-Z0-9])" 
    # Keyword + number (Accounts for OCR missing the space, e.g., "Chapter3")
    r"|(?:(?:Chapter|Section|Part|Appendix)\s*[\dIVXivx]+.?)" 
    r")$",
    re.MULTILINE | re.IGNORECASE, # Added ignorecase for 'chapter' vs 'Chapter'
)


def _detect_text_headers(text: str) -> list[tuple[str, str]]:
    """
    Returns [(header_label, section_body), ...] or [] if no headers found.
    Tries Markdown headers first, then generic pattern headers.
    """
    # --- Markdown-style ---
    matches = list(_MARKDOWN_HEADER_RE.finditer(text))
    if len(matches) >= 2:
        return _build_sections_from_matches(text, matches, label_group=2)

    # --- Generic / academic style ---
    matches = list(_GENERIC_HEADER_RE.finditer(text))
    if len(matches) >= 2:
        return _build_sections_from_matches(text, matches, label_group=0)

    return []


def _build_sections_from_matches(
    text: str,
    matches: list,
    label_group: int,
) -> list[tuple[str, str]]:
    sections = []
    for idx, match in enumerate(matches):
        header = match.group(label_group).strip()
        body_start = match.end()
        body_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        body = text[body_start:body_end].strip()
        sections.append((header, body))
    return sections


def _paragraph_sections(text: str, splitter: "SemanticAwareTextSplitter | None" =None, min_chars: int = 100) -> list[tuple[str, str]]:
    """
    Splits text into paragraph blocks, merging tiny paragraphs
    with the next one to avoid single-sentence micro-chunks.
    """
    raw_paras = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    merged: list[str] = []
    buffer = ""
    for para in raw_paras:
        buffer = (buffer + "\n\n" + para).strip() if buffer else para
        if len(buffer) >= min_chars:
            merged.append(buffer)
            buffer = ""
    if buffer:
        if merged:
            merged[-1] += "\n\n" + buffer   # attach trailing fragment to last chunk
        else:
            merged.append(buffer)

    if not splitter:
        return [("", block) for block in merged]

    # Sub-split oversized paragraphs semantically (no header label needed here)
    sections: list[tuple[str, str]] = []
    max_size = splitter.MAX_CHUNK_CHARS
    for block in merged:
        if len(block) > max_size:
            for sub in splitter.split(block):
                if sub.strip():
                    sections.append(("", sub.strip()))
        else:
            sections.append(("", block))
    return sections


# ---------------------------------------------------------------------------
# DOCX-specific section extractor
# ---------------------------------------------------------------------------
def _is_toc_page(text: str) -> bool:
        """
        Detects if a block of text is primarily a Table of Contents.
        Looks for lines ending in multiple dots followed by numbers.
        """
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if not lines:
            return False
            
        # Regex: looks for 4+ dots, optional spaces, and digits at the end of a line
        toc_line_pattern = re.compile(r'\.{4,}\s*\d+$')
        
        toc_count = sum(1 for line in lines if toc_line_pattern.search(line))
        ratio = toc_count / len(lines)
        
        # If more than 25% of the page is TOC lines, drop it
        return ratio > 0.25
def _extract_docx_sections(docx: DocxDocument) -> list[tuple[str, str]]:
    """
    Walk paragraphs; whenever we hit a Heading style, start a new section.
    Returns [(heading_text, section_body), ...].
    Falls back to returning all text as one unnamed section if no headings exist.
    """
    text_sections: list[tuple[str, str]] = []
    table_sections: list[tuple[str, str]] = []
    current_heading = ""
    current_body_lines: list[str] = []

    for para in docx.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if style_name.startswith("Heading"):
            if current_body_lines:
                text_sections.append((current_heading, "\n".join(current_body_lines)))
            current_heading = text
            current_body_lines = []
        else:
            current_body_lines.append(text)

    if current_body_lines:
        text_sections.append((current_heading, "\n".join(current_body_lines)))

    _semantic_splitter = SemanticAwareTextSplitter()

    for i, table in enumerate(docx.tables):
        data = []
        for row in table.rows:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            data.append(row_data)
            
        if len(data) > 1:
            df = pd.DataFrame(data[1:], columns=data[0])
            try:
                rows_per_chunk = _semantic_splitter.dynamic_table_rows(df)
                for start_row in range(0, len(df), rows_per_chunk):
                    df_chunk = df.iloc[start_row : start_row + rows_per_chunk]
                    md_table = df_chunk.to_markdown(index=False)
                    
                    if md_table:
                        end_row = start_row + len(df_chunk)
                        header  = (
                            f"{current_heading} — Table {i + 1} "
                            f"(rows {start_row + 1}–{end_row})"
                        )
                        table_sections.append((header, md_table))
            except Exception as e:
                print(f"Skipping malformed DOCX table: {e}")

    return text_sections, table_sections


# ---------------------------------------------------------------------------
# PDF-specific header heuristic (applied after text extraction)
# ---------------------------------------------------------------------------

def _extract_pdf_sections(text: str) -> list[tuple[str, str]]:
    """
    Runs header detection on extracted PDF text.
    Identical to the generic text path but kept separate for clarity.
    """
    return _detect_text_headers(text) or _paragraph_sections(text)


# ---------------------------------------------------------------------------
# Main loader class
# ---------------------------------------------------------------------------
con
class MultiFormatDocumentLoader:
    """
    Load and process various document formats using header/topic-aware chunking.
    Chunks align with document structure (sections, headings) rather than
    arbitrary character counts.
    """

    def __init__(self):
        self.supported_extensions = {
            ".pdf", ".docx", ".csv", ".xlsx", ".xls", ".txt",
            ".png", ".jpg", ".jpeg", ".bmp", ".tiff"
        }
        self.splitter = HeaderAwareTextSplitter()
        self._semantic_splitter = SemanticAwareTextSplitter()


    # ------------------------------------------------------------------
    # OCR helpers (unchanged from original)
    # ------------------------------------------------------------------
    def _detect_table_grid_cv2(self, image_path: str) -> bool:
        """
        Quickly scans an image for horizontal and vertical grid lines.
        Returns True if a table structure is detected.
        """
        try:
            img = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
            if img is None:
                return False

            # Binarize the image (invert so lines are white, background is black)
            _, thresh = cv2.threshold(img, 128, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)

            # Define kernels to detect horizontal and vertical lines
            # The length of the kernel dictates how long a line must be to be detected
            line_min_length = np.array(img).shape[1] // 10 
            
            kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (line_min_length, 1))
            kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, line_min_length))

            # Isolate horizontal lines
            horizontal = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel_h)
            # Isolate vertical lines
            vertical = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel_v)

            # Find intersections (where a horizontal and vertical line meet)
            intersections = cv2.bitwise_and(horizontal, vertical)

            # Count the intersections
            contours, _ = cv2.findContours(intersections, cv2.RETR_CCOMP, cv2.CHAIN_APPROX_SIMPLE)
            
            # If we have a decent number of intersections (e.g., 4+ corners), it's likely a table
            return len(contours) > 6

        except Exception as e:
            print(f"⚠️ Error during OpenCV table detection: {e}")
            return False
    def _looks_like_borderless_table(self, text: str) -> bool:
        """
        Analyzes OCR text for tabular spacing.
        Returns True if multiple lines have 3+ distinct columns separated by wide spaces.
        """
        if not text:
            return False
        lines = text.split("\n")
        tabular_line_count = sum(
            1 for line in lines if len(re.split(r"\s{3,}", line.strip())) >= 3
        )
        return tabular_line_count >= 3
    async def _get_ocr_text_async(self, file_path: str, origin_info: str = "Image") -> str:
        ocr_text = ""
        confidence = 0
        has_grid = await asyncio.to_thread(self._detect_table_grid_cv2, file_path)
        if not has_grid:
            try:
                def run_tesseract():
                    img = cv2.imread(file_path)
                    if img is not None:
                        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        ocr_data = pytesseract.image_to_data(thresh, output_type=pytesseract.Output.DICT)
                        ocr_text = pytesseract.image_to_string(thresh).strip()
                        confidences = [int(c) for c in ocr_data["conf"] if int(c) != -1]
                        confidence = sum(confidences) / len(confidences) if confidences else 0
                        return ocr_text, confidence
                ocr_text, confidence = await asyncio.to_thread(run_tesseract)
            except Exception as e:
                print(f"⚠️  Tesseract failed for {origin_info}: {e}")

        if has_grid or self._should_use_vlm_fallback(ocr_text, confidence):
            reason = "Grid detected" if has_grid else "Low quality/Borderless table detected"
            print(f"🔄 {reason} for {origin_info}. Falling back to VLM...")
            try:
                vlm_prompt = (
                    "Extract all readable text from this image. "
                    "CRITICAL INSTRUCTION: If the image contains a table, grid, "
                    "or structured data, you MUST format it as a valid Markdown "
                    "table using '|' to separate columns and '-' to separate the "
                    "header row. Preserve the exact row and column structure. "
                    "Return ONLY the markdown content, no conversational filler."
                )
                parts = []
                async for chunk in query_ollama(
                    prompt=vlm_prompt + origin_info,
                    model="codez-ocr:latest",
                    image_path=file_path,
                    keep_alive=0,
                    stream=False,
                ):
                    parts.append(chunk)
                ocr_text = "".join(parts).strip()
            except Exception as e:
                print(f"❌ VLM fallback failed: {e}")

        return ocr_text

    def _should_use_vlm_fallback(self, text: str, confidence: float) -> bool:
        if not text or len(text) < 10 or confidence < 60:
            return True
        if self._looks_like_borderless_table(text):
            return True
        words = text.split()
        if not words:
            return True
        return (sum(len(w) for w in words) / len(words)) < 2.5

    def _run_async_ocr(self, file_path: str, origin_info: str) -> str:
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
        return loop.run_until_complete(self._get_ocr_text_async(file_path, origin_info))

    # ------------------------------------------------------------------
    # Format loaders
    # ------------------------------------------------------------------

    def load_image(self, file_path: str) -> list[Document]:
        """
        Images produce a single document — OCR output has no reliable
        heading structure to split on.
        """
        print(f"⏳ Processing image: {os.path.basename(file_path)}")
        text = self._run_async_ocr(file_path, os.path.basename(file_path))
        if not text:
            return []
        # Still attempt header-based splitting in case the image is a scanned doc
        metadata = {"source": os.path.basename(file_path), "file_type": "image"}
        return self.splitter.split_plain_text(text, metadata)

    def load_pdf(self, file_path: str) -> list[Document]:
        if os.path.getsize(file_path) == 0:
            print(f"Skipping empty file: {file_path}")
            return []

        table_docs: list[Document] = []
        # NEW: collect (page_num, page_text) instead of one combined string
        page_texts: list[tuple[int, str]] = []

        base_metadata = {
            "source": os.path.basename(file_path),
            "file_type": "pdf",
        }

        try:
            with fitz.open(file_path) as f:
                for page_num, page in enumerate(f):
                    page_text = page.get_text().strip() or ""
                    if _is_toc_page(page_text):
                        print(f"⏩ Skipping TOC on page {page_num + 1}")
                        continue
                    tables = page.find_tables()
                    if tables:
                        page_text += "\n\n### Extracted Tabular Data\n\n"
                        for idx, tab in enumerate(tables):
                            try:
                                df=tab.to_pandas()
                                df.dropna(how="all", inplace=True)  # Drop empty rows
                                if not df.empty:
                                    rows_per_chunk = (self._semantic_splitter.dynamic_table_rows(df))
                                    for start_row in range(0, len(df), rows_per_chunk):
                                        df_chunk = df.iloc[start_row : start_row + rows_per_chunk]
                                        md_table = df_chunk.to_markdown(index=False)
                                        
                                        if md_table:
                                            end_row = start_row + len(df_chunk)
                                            table_meta = {
                                                **base_metadata,
                                                "page": page_num + 1,
                                            }
                                            header = (
                                                f"Table {idx + 1} "
                                                f"(page {page_num + 1}, "
                                                f"rows {start_row + 1}–{end_row})"
                                            )
                                            table_docs.extend(
                                                self.splitter.split_sections_with_context(
                                                    [(header, md_table)],
                                                    table_meta,
                                                    is_table=True,
                                                )
                                            )
                            except Exception as e:
                                print(f"Error extracting table from PDF page: {e}")

                    # ── Accumulate text with its page number ──────────────
                    if page_text:
                        page_texts.append((page_num + 1, page_text))   # 1-based

        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
            return []

        text_docs: list[Document] = []

        total_text = "".join(t for _, t in page_texts)

        if len(total_text.strip()) > 200:
            bridge = ""
            for page_num, page_text in page_texts:
                # prepend the last sentence from the previous page
                text_with_bridge = (bridge + " " + page_text).strip() if bridge else page_text

                page_meta = {**base_metadata, "page": page_num}
                sections = _extract_pdf_sections(text_with_bridge)
                text_docs.extend(
                    self.splitter.split_sections_with_context(
                        sections, page_meta, is_table=False
                    )
                )

                # carry the last sentence of this page into the next iteration
                sentences = page_text.split(". ")
                bridge = sentences[-1].strip() if sentences else ""
        else:
            # ── Scanned PDF: OCR path (unchanged logic, already per-page) ──
            print(f"⚠️  Scanned PDF detected: {os.path.basename(file_path)}")
            try:
                images = convert_from_path(file_path)
                for i, img in enumerate(images):
                    temp_path = f"temp_{os.getpid()}_{i}.jpg"
                    img = img.convert("RGB")
                    img.save(temp_path, "JPEG", quality=90)
                    ocr_text = self._run_async_ocr(temp_path, f"Page {i + 1}")
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
                    if not ocr_text:
                        continue
                    page_meta = {**base_metadata, "file_type": "pdf_scanned", "page": i + 1}
                    sections = _detect_text_headers(ocr_text) or [("", ocr_text)]
                    text_docs.extend(
                        self.splitter.split_sections_with_context(sections, page_meta)
                    )
            except Exception as e:
                print(f"Error loading TXT {file_path}: {e}")
                return []

        final_docs =table_docs + text_docs
        if not final_docs and total_text.strip():
            return self.splitter.split_plain_text(total_text, base_metadata)
        return final_docs
    
    def load_docx(self, file_path: str) -> list[Document]:
        try:
            docx = DocxDocument(file_path)
        except Exception as e:
            print(f"Error loading DOCX {file_path}: {e}")
            return []

        text_sections, table_sections = _extract_docx_sections(docx)
        metadata = {"source": os.path.basename(file_path), "file_type": "docx"}
        
        # Process text and tables separately to protect the tables
        docs = self.splitter.split_sections_with_context(text_sections, metadata, is_table=False)
        docs.extend(self.splitter.split_sections_with_context(table_sections, metadata, is_table=True))
        return docs

    def load_csv(self, file_path: str, rows_per_chunk: int = 250) -> list[Document]:
        """
        CSVs have no heading structure; row-group chunking is the right
        semantic boundary here — unchanged from original.
        """
        docs = []
        try:
            for i, df_chunk in enumerate(
                pd.read_csv(file_path, on_bad_lines="skip", chunksize=rows_per_chunk, low_memory=True)
            ):
                buf = io.StringIO()
                content=df_chunk.to_markdown(index=False)
                if content:
                    start = i * rows_per_chunk + 1
                    docs.append(Document(
                        page_content=content,
                        metadata={
                            "source": os.path.basename(file_path),
                            "rows": f"{start}-{start + len(df_chunk) - 1}",
                            "file_type": "csv_chunk",
                        },
                    ))
        except Exception as e:
            print(f"Error loading CSV {file_path}: {e}")
        return docs

    def load_excel(self, file_path: str, rows_per_chunk: int = 250) -> list[Document]:
        """
        Each sheet is treated as a named section; within a sheet,
        further splitting uses the plain-text path if the sheet is large.
        """
        docs = []
        try:
            xls = pd.ExcelFile(file_path)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Manually slice the dataframe into chunks
                for start_row in range(0, len(df), rows_per_chunk):
                    df_chunk = df.iloc[start_row : start_row + rows_per_chunk]
                    content = df_chunk.to_markdown(index=False)
                    
                    if content:
                        start_idx = start_row + 1
                        end_idx = start_row + len(df_chunk)
                        docs.append(Document(
                            page_content=content,
                            metadata={
                                "source": os.path.basename(file_path),
                                "sheet": sheet_name,
                                "rows": f"{start_idx}-{end_idx}",
                                "file_type": "excel_chunk",
                                "is_table": True
                            },
                        ))
        except Exception as e:
            print(f"Error loading Excel {file_path}: {e}")
        return docs

    def load_txt(self, file_path: str) -> list[Document]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except Exception as e:
            print(f"Error loading TXT {file_path}: {e}")
            return []

        metadata = {"source": os.path.basename(file_path), "file_type": "txt"}
        return self.splitter.split_plain_text(text, metadata)

    # ------------------------------------------------------------------
    # Dispatcher
    # ------------------------------------------------------------------

    def load_document(self, file_path: str) -> list[Document]:
        ext = Path(file_path).suffix.lower()
        dispatch = {
            ".pdf":  self.load_pdf,
            ".docx": self.load_docx,
            ".csv":  self.load_csv,
            ".xlsx": self.load_excel,
            ".xls":  self.load_excel,
            ".txt":  self.load_txt,
            ".png":  self.load_image,
            ".jpg":  self.load_image,
            ".jpeg": self.load_image,
            ".bmp":  self.load_image,
            ".tiff": self.load_image,
        }
        loader = dispatch.get(ext)
        if loader:
            return loader(file_path)
        print(f"Unsupported file type: {ext}")
        return []
    
    
    
def dump_chunks_to_file(
    docs: list,
    output_path: str = "vector.txt",
    encoding: str = "utf-8",
    no_change: bool = False,
    mode: str = "w"  # <-- Added mode to allow appending
) -> None:
    """
    Write document chunks along with metadata to a plain-text file.
    Supports streaming chunks in real-time using mode="a".
    """
    if no_change or not docs:
        return

    KNOWN_KEYS = ["chunk_id", "source", "file_type", "section", "page", "rows", "sheet"]
    LABELS = {
        "chunk_id":  "CHUNK ID ",
        "source":    "SOURCE   ",
        "file_type": "FILE TYPE",
        "section":   "SECTION  ",
        "page":      "PAGE     ",
        "rows":      "ROWS     ",
        "sheet":     "SHEET    ",
    }

    WIDE  = "=" * 80
    THIN  = "-" * 80

    with open(output_path, mode, encoding=encoding, errors="replace") as f:
        # Only write a main header if we are starting a fresh file
        if mode == "w":
            f.write("=== VECTOR CHUNK DUMP ===\n\n")

        for doc in docs:
            meta = doc.metadata or {}
            content = doc.page_content or ""

            f.write(WIDE + "\n")
            f.write("CHUNK\n")
            f.write(THIN + "\n")

            # Known metadata fields
            for key in KNOWN_KEYS:
                if key in meta:
                    label = LABELS[key]
                    f.write(f"{label}: {meta[key]}\n")

            # Char count
            f.write(f"CHARS    : {len(content)}\n")

            # Any extra/unknown metadata keys
            extra = {k: v for k, v in meta.items() if k not in KNOWN_KEYS}
            if extra:
                f.write(f"OTHER    : {extra}\n")

            f.write(THIN + "\n")
            f.write(content)
            f.write("\n" + WIDE + "\n\n")